'''
Created on 12/10/2018

@author: Jiaming
'''
import docx
import re

test_entity = {"Type": "Object"}
test_lex_dic = {"Pred": ["bowl"],
                "Type": ["physobj*artifact"]}
test_typ_dic = {"Head": ["cylindroid[1]"],
           "Components": [["Component", "Value"], ["surface[1]"], ["interior[2]"]],
           "Concavity": ["Concave[2]"],
           "RotatSym": ["Y"],
           "ReflSym": ["XY, YZ"]}
test_habt_dic = {"Intrinsic": [["Intr", "Name", "Value"], ["UP[3]", "align(Y,E_y)"], ["TOP[3]", "top(+Y)"]],
            "Extrinsic": [["Extr", "Name", "Value"], ["UP[4]", "align(Y,E_y)"], ["TOP[4]", "top(-Y)"]]}
test_afford_dic = {"Affordances": [["Affordance", "Formula"],
                              ["H->[put(x, on([1]))]support([1], x)"],
                              ["H->[put(x, in([2]))]contain([2], x)"],
                              ["H[4]->[put([2], on(x))]contain([2], x)"],
                              ["H->[lift(x, [1])]hold(x, [1])"],
                              ["H->[grasp(x, [1])]"]]}
test_embd_dic = {"Scale": ["<agent"],
            "Movable": ["true"]}

entity = {}
lex_dic = {}
typ_dic = {}
habt_dic = {}
afford_dic = {}
embd_dic = {}

lex_text = ""
typ_text = ""
habt_text = ""
afford_text = ""
embd_text = ""

'''
def get_full_text(filename):

    doc = docx.Document(filename)
    text_list = []
    for paragraph in doc.paragraphs:
        text_list.append(paragraph.text)
    return '\n'.join(text_list)
'''
shape_inventory = {"paraboloid", "sheet", "ellipsoid", "cylinder", "prism", "rectangular", "slab", "elliptic"}
axis_inventory = {"x-axis", "y-axis", "z-axis", "xy-axis", "yz-axis", "xz-axis",
                  "x axis", "y axis", "z axis", "xy axis", "yz axis", "xz axis"}

def get_runs(filename):

    doc = docx.Document(filename)
    runs_list = []
    new_runs = []
    count = 0

    for paragraph in doc.paragraphs:
        for x in range(0, len(paragraph.runs)):
            if paragraph.runs[x].underline is True:
                count += 1
                runs_list.append([paragraph.runs[x].text])
            else:
                runs_list[count-1].append(paragraph.runs[x].text)

    for run in runs_list:
        new_runs.append(' '.join(run))

    return new_runs


def lex_process(text):
    pred_pattern = re.compile(r":.*")
    pred_match = pred_pattern.findall(text)
    if pred_match is not None:
        lex_dic.update({"Pred": [pred_match[0].lower().split()[-1]]})
        # need more clue to identify artifact object
        lex_dic.update({"Type": ["physobj"]})


def tpy_process(text):

    head_pattern = re.compile(r"shape of.*?\.")
    head_match = head_pattern.findall(text)
    if head_match is not None:
        shape = [x for x in shape_inventory if x in str(head_match).lower()]
        typ_dic.update({"Head": [" ".join(shape)]})

    comp_pattern = re.compile(r"Components:.*")
    comp_match = comp_pattern.findall(text)
    if comp_match is not None:
        comp_lst = [["Component", "Value"]]
        for comp in comp_match[0].lower().split()[1:]:
            comp_lst.append([comp])
        typ_dic.update({"Components": comp_lst})

    concv_pattern = re.compile(r"(inside.*?\.|scoop.*?\.)")
    concv_match = concv_pattern.findall(text)
    if concv_match is not None:
        typ_dic.update({"Concavity": ["Concave"]})
    elif re.compile(r"flat.*?\.").findall(text) is not None:
        typ_dic.update({"Concavity": ["Flat"]})
    else:
        typ_dic.update({"Concavity": ["Convex"]})

    rotat_pattern = re.compile(r"rotat.*?\.")
    rotat_match = rotat_pattern.findall(text)
    if rotat_match is not None:
        rotat = [x for x in axis_inventory if x in str(rotat_match).lower()]
        typ_dic.update({"RotatSym": [" ".join(rotat)]})

    refl_pattern = re.compile(r"refl.*?\.")
    refl_match = refl_pattern.findall(text)
    if refl_match is not None:
        refl = [x for x in axis_inventory if x in str(refl_match).lower()]
        typ_dic.update({"ReflSym": [" ".join(refl)]})


def habt_process(text):
    habt_pattern = re.compile(r"can.*?\.")
    habt_match = habt_pattern.findall(text)
    if habt_match is not None:
        habt_lst = [["Intr", "Value"]]
        for habt in habt_match:
            habt_lst.append([habt])
    habt_dic.update({"Intrinsic": habt_lst})


def afford_process(text):
    affd_pattern = re.compile(r"\bIf.*?\.")
    affd_match = affd_pattern.findall(text)
    if affd_match is not None:
        affd_lst = [["Affordance", "Formula"]]
        for affd in affd_match:
            affd_lst.append([affd])
    afford_dic.update({"Affordances": affd_lst})


# full_text = re.compile(r"\n").sub(" ", get_full_text("Bowl_a.docx"))

for paragraph in get_runs("bowl_annotation.docx"):
    if re.match(r"\bObject", paragraph) is not None:
        lex_text = paragraph
        entity = {"Type": "Object"}
    if re.match(r"\bComponent", paragraph) is not None:
        typ_text += paragraph
    if re.match(r"\bDescription", paragraph) is not None:
        typ_text += paragraph
    if re.match(r"\bState", paragraph) is not None:
        habt_text += paragraph
        afford_text += paragraph
    if re.match(r"\b(Action|Activities)", paragraph) is not None:
        habt_text += paragraph
        afford_text += paragraph

lex_process(lex_text)
tpy_process(typ_text)
habt_process(habt_text)
afford_process(afford_text)


'''print(entity)
print(lex_dic)
print(typ_dic)
print(afford_dic)
print(entity)
print(embd_dic)


print("Lex paragraph: "+lex_text)
print("type paragraph: "+typ_text)
print("habitate paragraph: "+habt_text)
print("affordance paragraph: "+afford_text)'''

